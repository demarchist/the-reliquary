"""
Convert campaign diary .docx files and lore document into structured Markdown
for The Reliquary Astro content collections.

Usage:
    python scripts/convert-docx.py

Expects source files in ~/Downloads/:
    - RoA_ Campaign Diary.docx       (Chapters 1-8)
    - ROA_ Campaign Diary pt. 2.docx (Chapters 8 ctd - 14 start)
    - ROA_ Campaign Diary - Chapter 14.docx
    - D&D Lore.docx
"""

import os
import re
import shutil
from pathlib import Path

import docx

# Paths
DOWNLOADS = Path(os.path.expanduser("~/Downloads"))
PROJECT_ROOT = Path(__file__).resolve().parent.parent
DIARY_DIR = PROJECT_ROOT / "src" / "content" / "diary"
LORE_DIR = PROJECT_ROOT / "src" / "content" / "lore"

SOURCE_FILES = [
    DOWNLOADS / "RoA_ Campaign Diary.docx",
    DOWNLOADS / "ROA_ Campaign Diary pt. 2.docx",
    DOWNLOADS / "ROA_ Campaign Diary - Chapter 14.docx",
]

LORE_FILE = DOWNLOADS / "D&D Lore.docx"


def slugify(text: str) -> str:
    """Convert text to a URL-friendly slug."""
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text)
    return text.strip("-")


def para_to_markdown(para) -> str:
    """Convert a docx paragraph to markdown, preserving bold/italic."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def is_chapter_heading(para) -> bool:
    """Check if paragraph is a chapter title."""
    style = para.style.name.lower() if para.style else ""
    text = para.text.strip().lower()
    return style == "title" and "chapter" in text


def is_game_heading(para) -> bool:
    """Check if paragraph is a game/session heading."""
    style = para.style.name.lower() if para.style else ""
    text = para.text.strip()
    if not text:
        return False
    # Match "Game N" patterns - either as Heading 1 or as short standalone text
    game_pattern = re.match(r"^Game\s+\d", text, re.IGNORECASE)
    if style in ("heading 1", "title") and game_pattern:
        return True
    # Also match plain "Game N" lines that are short (not body text)
    if game_pattern and len(text) < 80:
        return True
    return False


def parse_chapter_info(text: str) -> tuple[int, str]:
    """Extract chapter number and title from heading text."""
    # Match patterns like "Chapter 1, Phandalin: Game 1" or "Chapter 9: The Demon Lords"
    m = re.match(
        r"Chapter\s+(\d+)[,.:]\s*(.*?)(?:\s*:\s*Game\s+\d+)?$", text, re.IGNORECASE
    )
    if m:
        return int(m.group(1)), m.group(2).strip().rstrip(":")
    # Fallback: just get the number
    m = re.match(r"Chapter\s+(\d+)", text, re.IGNORECASE)
    if m:
        return int(m.group(1)), ""
    return 0, text


def parse_game_number(text: str) -> tuple[int, str]:
    """Extract game number and optional subtitle from heading text."""
    # Match "Game 23 - The Orc and the Alchemist" or "Game 4,5,6,7,8 - text"
    m = re.match(r"Game\s+([\d,]+)\s*[-:]?\s*(.*)", text, re.IGNORECASE)
    if m:
        # Take the first number for multi-game entries
        nums = m.group(1).split(",")
        return int(nums[0].strip()), m.group(2).strip()
    m = re.match(r"Game\s+(\d+)", text, re.IGNORECASE)
    if m:
        return int(m.group(1)), ""
    return 0, text


def process_diary_file(filepath: Path, state: dict):
    """Process a single diary docx file, updating running state."""
    print(f"Processing: {filepath.name}")
    doc = docx.Document(str(filepath))

    current_chapter = state.get("chapter", 0)
    current_chapter_title = state.get("chapter_title", "")
    current_game = state.get("game", 0)
    current_game_title = state.get("game_title", "")
    paragraphs: list[str] = state.get("pending_paragraphs", [])

    def flush_session():
        """Write the current session to a file."""
        nonlocal paragraphs
        if current_chapter == 0 or current_game == 0 or not paragraphs:
            paragraphs = []
            return

        # Clean up paragraphs
        content = "\n\n".join(p for p in paragraphs if p.strip())

        game_title = current_game_title or f"Session {current_game}"
        title = f"{current_chapter_title} - {game_title}" if current_chapter_title else game_title

        filename = f"chapter-{current_chapter:02d}-session-{current_game:02d}.md"
        filepath_out = DIARY_DIR / filename

        frontmatter = f"""---
title: "{title}"
chapter: {current_chapter}
chapterTitle: "{current_chapter_title}"
session: {current_game}
---"""

        filepath_out.write_text(f"{frontmatter}\n\n{content}\n", encoding="utf-8")
        print(f"  Wrote: {filename}")
        paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()

        if is_chapter_heading(para):
            flush_session()
            current_chapter, current_chapter_title = parse_chapter_info(text)
            current_game = 0
            current_game_title = ""
            # The chapter heading often includes "Game 1"
            if re.search(r"Game\s+\d", text, re.IGNORECASE):
                gm = re.search(r"Game\s+(\d+)", text, re.IGNORECASE)
                if gm:
                    current_game = int(gm.group(1))
                    current_game_title = ""
            continue

        if is_game_heading(para):
            flush_session()
            current_game, current_game_title = parse_game_number(text)
            continue

        # Regular paragraph - convert to markdown
        md = para_to_markdown(para)
        if md.strip():
            paragraphs.append(md)

    # Update state for next file (don't flush - might continue in next file)
    state["chapter"] = current_chapter
    state["chapter_title"] = current_chapter_title
    state["game"] = current_game
    state["game_title"] = current_game_title
    state["pending_paragraphs"] = paragraphs


def process_lore_file(filepath: Path):
    """Convert the lore document to markdown."""
    print(f"Processing lore: {filepath.name}")
    doc = docx.Document(str(filepath))

    paragraphs = []
    for para in doc.paragraphs:
        md = para_to_markdown(para)
        if md.strip():
            paragraphs.append(md)

    content = "\n\n".join(paragraphs)
    out = LORE_DIR / "world-lore.md"
    out.write_text(f"# World Lore\n\n{content}\n", encoding="utf-8")
    print(f"  Wrote: world-lore.md")


def main():
    # Clean and recreate output directories
    if DIARY_DIR.exists():
        for f in DIARY_DIR.glob("*.md"):
            f.unlink()
    DIARY_DIR.mkdir(parents=True, exist_ok=True)

    if LORE_DIR.exists():
        for f in LORE_DIR.glob("*.md"):
            f.unlink()
    LORE_DIR.mkdir(parents=True, exist_ok=True)

    # Check source files exist
    missing = [f for f in SOURCE_FILES if not f.exists()]
    if missing:
        print("Missing source files:")
        for f in missing:
            print(f"  {f}")
        return

    if not LORE_FILE.exists():
        print(f"Missing lore file: {LORE_FILE}")
        return

    # Process diary files in order with shared state
    state: dict = {}

    for filepath in SOURCE_FILES:
        process_diary_file(filepath, state)

    # Flush any remaining content from the last file
    if state.get("pending_paragraphs") and state.get("chapter") and state.get("game"):
        ch = state["chapter"]
        gm = state["game"]
        ch_title = state.get("chapter_title", "")
        gm_title = state.get("game_title", f"Session {gm}")
        title = f"{ch_title} - {gm_title}" if ch_title else gm_title

        filename = f"chapter-{ch:02d}-session-{gm:02d}.md"
        content = "\n\n".join(
            p for p in state["pending_paragraphs"] if p.strip()
        )
        frontmatter = f"""---
title: "{title}"
chapter: {ch}
chapterTitle: "{ch_title}"
session: {gm}
---"""
        (DIARY_DIR / filename).write_text(
            f"{frontmatter}\n\n{content}\n", encoding="utf-8"
        )
        print(f"  Wrote: {filename}")

    # Process lore
    process_lore_file(LORE_FILE)

    # Count results
    diary_count = len(list(DIARY_DIR.glob("*.md")))
    print(f"\nDone! Generated {diary_count} diary entries and 1 lore file.")


if __name__ == "__main__":
    main()
